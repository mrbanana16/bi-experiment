# 预期实现：输出单个回答或完整对话记录为markdown（直接存储json message），word（使用pandoc）或者pdf（docx2pdf）

# PDF 导出当前依赖 docx2pdf，运行机器需要安装 Microsoft Word
import argparse
import json
import platform
import re
import sys
import tempfile
from datetime import datetime
from html import escape as escape_html
from pathlib import Path
import pypandoc
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[2]
DEFAULT_HISTORY_PATH = PROJECT_ROOT / "result" / "ai" / "ai-history" / "history_20260613_203717_922994.json"
DEFAULT_OUTPUT_DIR = PROJECT_ROOT / "result" / "ai" / "ai-output"

ROLE_LABELS = {
    "user": "用户",
    "assistant": "AI 助手",
    "system": "系统",
}
DOCUMENT_FONT_NAME = "DengXian"
DOCUMENT_COLOR = "000000"
EXPORT_FORMAT_ALIASES = {
    "markdown": "markdown",
    "md": "markdown",
    "docx": "docx",
    "pdf": "pdf",
}


def read_history(history_path):
    with history_path.open("r", encoding="utf-8") as file:
        history = json.load(file)

    if not isinstance(history, dict):
        raise ValueError("history json root must be an object")

    messages = history.get("messages")
    if not isinstance(messages, list):
        raise ValueError("history json must contain a messages list")

    return history


def export_full_history(history_path=DEFAULT_HISTORY_PATH, output_dir=DEFAULT_OUTPUT_DIR, base_name=None, file_format="markdown", log=print):
    history_path = Path(history_path).resolve()
    output_dir = Path(output_dir).resolve()
    export_format = normalize_export_format(file_format)
    log(f"读取历史对话：{history_path}")
    history = read_history(history_path)

    output_dir.mkdir(parents=True, exist_ok=True)
    base_name = base_name or f"{history.get('id') or history_path.stem}_conversation"
    log(f"导出目录：{output_dir}")
    log(f"导出文件名：{base_name}")
    log(f"导出格式：{export_format}")

    md_path = output_dir / f"{base_name}.md"
    docx_path = output_dir / f"{base_name}.docx"
    pdf_path = output_dir / f"{base_name}.pdf"

    markdown_content = build_export_markdown(history)

    if export_format == "markdown":
        log(f"生成 Markdown：{md_path}")
        md_path.write_text(markdown_content, encoding="utf-8")
        paths = {"markdown": md_path}
    elif export_format == "docx":
        log(f"生成 Word 文档：{docx_path}")
        convert_markdown_content_to_docx(markdown_content, docx_path)
        paths = {"docx": docx_path}
    else:
        log(f"生成 PDF：{pdf_path}")
        with tempfile.TemporaryDirectory(prefix="ai-export-") as temp_dir:
            temp_docx_path = Path(temp_dir) / f"{base_name}.docx"
            convert_markdown_content_to_docx(markdown_content, temp_docx_path)
            convert_docx_to_pdf(temp_docx_path, pdf_path, log=log)
        paths = {"pdf": pdf_path}

    log("导出完成")

    return paths


def normalize_export_format(file_format):
    export_format = EXPORT_FORMAT_ALIASES.get(str(file_format or "").lower())
    if not export_format:
        raise ValueError("unsupported export format")
    return export_format


def build_export_markdown(history):
    if history.get("export_kind") == "report":
        return build_report_markdown(history)
    return build_history_markdown(history)


def build_report_markdown(history):
    report_message = next(
        (message for message in reversed(history.get("messages") or []) if isinstance(message, dict) and message.get("role") == "assistant"),
        {},
    )
    model = report_message.get("model") or history.get("model") or "AI 模型"
    content = normalize_markdown_content(normalize_text(report_message.get("content")))
    generated_at = history.get("updated_at") or datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    footer = [
        "",
        "---",
        "",
        f"*以上内容由{model}生成，请核查是否正确。",
        f"生成时间：{generated_at}",
    ]
    return "\n".join([content.rstrip(), *footer]).rstrip() + "\n"


def build_history_markdown(history):
    history_id = history.get("id") or "未命名对话"
    created_at = history.get("created_at") or ""
    updated_at = history.get("updated_at") or ""
    model = history.get("model") or ""
    selected_results = history.get("selected_results") or []
    messages = history.get("messages") or []

    lines = [
        f"# AI 对话记录：{history_id}",
        "",
        "## 基本信息",
        "",
        f"- 对话 ID：{history_id}",
        f"- 创建时间：{created_at or '未知'}",
        f"- 更新时间：{updated_at or '未知'}",
        f"- 模型：{model or '未知'}",
        f"- 导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
    ]

    if selected_results:
        lines.extend(build_selected_results_markdown(selected_results))

    lines.extend([
        "## 对话内容",
        "",
    ])

    for round_index, round_messages in enumerate(build_conversation_rounds(messages), start=1):
        lines.extend(build_round_markdown(round_index, round_messages))

    return "\n".join(lines).rstrip() + "\n"


def build_selected_results_markdown(selected_results):
    lines = [
        "## 已选分析结果",
        "",
        "| 序号 | 类型 | 文件名 | 路径 |",
        "| --- | --- | --- | --- |",
    ]

    for index, item in enumerate(selected_results, start=1):
        if not isinstance(item, dict):
            continue

        result_type = escape_table_text(item.get("type") or item.get("category") or "")
        name = escape_table_text(item.get("name") or "")
        path = escape_table_text(item.get("path") or "")
        lines.append(f"| {index} | {result_type} | {name} | `{path}` |")

    lines.append("")
    return lines


def build_conversation_rounds(messages):
    rounds = []
    current_round = None

    for message in messages:
        if not isinstance(message, dict):
            if current_round is None:
                current_round = []
                rounds.append(current_round)
            current_round.append(message)
            continue

        if message.get("role") == "user" or current_round is None:
            current_round = []
            rounds.append(current_round)

        current_round.append(message)

    return rounds


def build_round_markdown(round_index, messages):
    lines = [
        f"### 第 {round_index} 轮",
        "",
    ]

    for message in messages:
        lines.extend(build_message_markdown(message))

    lines.append("---")
    lines.append("")
    return lines


def build_message_markdown(message):
    if not isinstance(message, dict):
        return [
            "#### 未知消息",
            "",
            str(message),
            "",
        ]

    role = message.get("role") or "unknown"
    role_label = ROLE_LABELS.get(role, role)
    created_at = message.get("created_at") or ""
    model = message.get("model") or ""
    is_error = message.get("is_error")
    content = normalize_markdown_content(normalize_text(message.get("content")))
    message_summary = build_message_summary(message)

    role_heading = f"#### *{role_label}*"
    if message_summary:
        role_heading += f' <span style="font-size:8pt"><em>{message_summary}</em></span>'

    lines = [
        role_heading,
        "",
    ]
    lines.append(content or "_（空内容）_")
    lines.append("")
    return lines


def build_message_summary(message):
    summary_items = []
    created_at = message.get("created_at") or ""
    model = message.get("model") or ""
    is_error = message.get("is_error")

    if created_at:
        summary_items.append(created_at)
    if model:
        summary_items.append(model)
    if is_error:
        summary_items.append("错误")

    return escape_html(" / ".join(summary_items))


def normalize_text(value):
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    return json.dumps(value, ensure_ascii=False, indent=2)


def normalize_markdown_content(content):
    if not content:
        return ""

    lines = content.splitlines()
    normalized_lines = []

    for line in lines:
        stripped = line.lstrip()

        if normalized_lines and is_markdown_table_row(stripped):
            previous = normalized_lines[-1].strip()
            if previous and not is_markdown_table_row(previous):
                normalized_lines.append("")

        if normalized_lines and is_markdown_list_marker(stripped):
            previous = normalized_lines[-1].strip()
            if previous and not is_markdown_list_marker(previous) and not previous.startswith("|"):
                normalized_lines.append("")

        normalized_lines.append(line)

    return ensure_blank_lines_around_tables(normalized_lines)


def is_markdown_list_marker(text):
    return bool(re.match(r"^([-*+]|\d+\.)\s+", text))


def is_markdown_table_row(text):
    stripped = text.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def ensure_blank_lines_around_tables(lines):
    normalized = []
    index = 0

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if is_markdown_table_row(stripped):
            if normalized and normalized[-1].strip():
                normalized.append("")

            while index < len(lines) and is_markdown_table_row(lines[index].strip()):
                normalized.append(lines[index])
                index += 1

            if index < len(lines) and lines[index].strip():
                normalized.append("")
            continue

        normalized.append(line)
        index += 1

    return "\n".join(normalized).strip()


def escape_table_text(value):
    return str(value).replace("|", "\\|").replace("\n", " ")


def convert_markdown_to_docx(md_path, docx_path):
    pypandoc.convert_file(
        str(md_path),
        "docx",
        outputfile=str(docx_path),
        extra_args=["--standalone"],
    )
    polish_docx(docx_path)


def convert_markdown_content_to_docx(markdown_content, docx_path):
    docx_path = Path(docx_path)
    with tempfile.NamedTemporaryFile("w", suffix=".md", encoding="utf-8", delete=False) as temp_file:
        temp_file.write(markdown_content)
        temp_md_path = Path(temp_file.name)

    try:
        convert_markdown_to_docx(temp_md_path, docx_path)
    finally:
        temp_md_path.unlink(missing_ok=True)


def polish_docx(docx_path):
    document = Document(docx_path)
    polish_document_styles(document)

    for paragraph in iter_document_paragraphs(document):
        if paragraph.style and paragraph.style.name == "Heading 4":
            polish_message_heading(paragraph)
        else:
            polish_paragraph_runs(paragraph)

    document.save(docx_path)


def polish_document_styles(document):
    style_specs = {
        "Normal": {"size": 10.5, "bold": False},
        "Body Text": {"size": 10.5, "bold": False},
        "First Paragraph": {"size": 10.5, "bold": False},
        "Compact": {"size": 10.5, "bold": False},
        "Table": {"size": 10.5, "bold": False},
        "Heading 1": {"size": 18, "bold": True},
        "Heading 2": {"size": 16, "bold": True},
        "Heading 3": {"size": 12, "bold": True},
        "Heading 4": {"size": 12, "bold": True},
    }

    for style_name, spec in style_specs.items():
        if style_name not in document.styles:
            continue

        apply_font(document.styles[style_name].font, spec["size"], bold=spec["bold"])


def iter_document_paragraphs(document):
    for paragraph in document.paragraphs:
        yield paragraph

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def polish_message_heading(paragraph):
    text = paragraph.text.strip()
    role_label = find_role_label(text)

    if not role_label:
        return

    message_summary = text[len(role_label):].strip()
    paragraph.clear()

    role_run = paragraph.add_run(role_label)
    role_run.italic = True
    apply_font(role_run.font, 12)

    if message_summary:
        summary_run = paragraph.add_run(f" {message_summary}")
        summary_run.italic = True
        apply_font(summary_run.font, 8)


def polish_paragraph_runs(paragraph):
    for run in paragraph.runs:
        apply_font(run.font, size_for_paragraph(paragraph), bold=bold_for_paragraph(paragraph))


def size_for_paragraph(paragraph):
    if not paragraph.style:
        return 10.5

    style_name = paragraph.style.name
    if style_name == "Heading 1":
        return 18
    if style_name == "Heading 2":
        return 16
    if style_name in {"Heading 3", "Heading 4"}:
        return 12
    return 10.5


def bold_for_paragraph(paragraph):
    return bool(paragraph.style and paragraph.style.name in {"Heading 1", "Heading 2", "Heading 3", "Heading 4"})


def apply_font(font, size, bold=None):
    font.name = DOCUMENT_FONT_NAME
    font.size = Pt(size)
    font.color.rgb = RGBColor.from_string(DOCUMENT_COLOR)

    if bold is not None:
        font.bold = bold

    if font.element.rPr is not None:
        font.element.rPr.rFonts.set(qn("w:eastAsia"), DOCUMENT_FONT_NAME)


def find_role_label(text):
    for role_label in sorted(set(ROLE_LABELS.values()), key=len, reverse=True):
        if text.startswith(role_label):
            return role_label

    return None


def convert_docx_to_pdf(docx_path, pdf_path, log=print):
    word_error_message = "PDF 导出失败：未检测到可用的 Microsoft Word，或当前环境无法打开 Microsoft Word。请确认本机已安装 Microsoft Word 后重试。"

    try:
        if platform.system() == "Windows":
            convert_docx_to_pdf_with_windows_word(docx_path, pdf_path)
        else:
            convert_docx_to_pdf_with_docx2pdf(docx_path, pdf_path)
    except Exception as error:
        log(word_error_message)
        raise RuntimeError(f"{word_error_message} 原始错误：{error}") from error

    if not pdf_path.exists():
        log(word_error_message)
        raise RuntimeError(f"{word_error_message} 未生成 PDF 文件：{pdf_path}")


def convert_docx_to_pdf_with_docx2pdf(docx_path, pdf_path):
    from docx2pdf import convert as docx2pdf_convert

    docx2pdf_convert(str(docx_path), str(pdf_path))


def convert_docx_to_pdf_with_windows_word(docx_path, pdf_path):
    import pythoncom
    import win32com.client

    docx_path = str(Path(docx_path).resolve())
    pdf_path = str(Path(pdf_path).resolve())
    word = None
    document = None
    pythoncom.CoInitialize()

    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(docx_path, ReadOnly=True)
        document.SaveAs(pdf_path, FileFormat=17)
    finally:
        if document is not None:
            document.Close(False)
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()


def parse_args():
    parser = argparse.ArgumentParser(description="导出 AI 完整对话记录为指定格式。")
    parser.add_argument("--history", default=str(DEFAULT_HISTORY_PATH), help="历史对话 JSON 文件路径。")
    parser.add_argument("--output-dir", default=str(DEFAULT_OUTPUT_DIR), help="导出文件目录。")
    parser.add_argument("--base-name", default=None, help="导出文件名，不包含扩展名。")
    parser.add_argument("--format", default="markdown", choices=sorted(EXPORT_FORMAT_ALIASES), help="导出文件格式。")
    return parser.parse_args()


def main():
    args = parse_args()

    try:
        paths = export_full_history(args.history, args.output_dir, args.base_name, args.format)
    except Exception as error:
        print(f"导出失败：{error}", file=sys.stderr)
        return 1

    print("导出完成：")
    for file_type, file_path in paths.items():
        print(f"- {file_type}: {file_path}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
